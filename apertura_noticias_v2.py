#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Apertura de Noticias (HOY, prioriza fuentes; fallback controlado) + DOCX con síntesis
Autor: Mr. PC + Jazz

Instalación (en consola de Spyder):
    !python -m pip install requests feedparser pandas python-docx pytz newspaper3k lxml
"""

import os, time, re, requests, pytz, feedparser, pandas as pd
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime
from urllib.parse import quote_plus, urlparse, parse_qs
from pathlib import Path

# Word
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ====== EXTRACTORES DE TEXTO ======
_HAVE_NEWSPAPER = False
_HAVE_READABILITY = False
try:
    from newspaper import Article
    _HAVE_NEWSPAPER = True
except Exception:
    pass

try:
    from readability import Document as ReadabilityDoc
    _HAVE_READABILITY = True
except Exception:
    pass

# =========================
# CONFIGURACIÓN
# =========================
TZ = "America/Mexico_City"
OUT_DIR_ROOT = os.path.join(os.path.expanduser("~"), "aperturas")
NOMBRE_BASE = "apertura_noticias"
NOMBRE_CSV_LOG = "apertura_noticias_log.csv"

MAX_POR_TEMA = 2
REINTENTOS = 3
ESPERA_SEG = 1.2
HTTP_TIMEOUT = 25

# --- Email ---
ASUNTO_BASE = "Apertura de noticias"
EMAIL_TO = ["alvarocandia007@gmail.com"]
EMAIL_CC = []
GMAIL_USER = os.getenv("GMAIL_USER")
GMAIL_APP_PASS = os.getenv("GMAIL_APP_PASS")
CORTE_LABEL = os.getenv("CORTE_LABEL")  # ej. "Corte Matutino"

# Fuentes preferidas
FUENTES_PREFERIDAS = {
    "elfinanciero.com.mx",
    "eleconomista.com.mx",
    "jornada.com.mx",
    "eluniversal.com.mx",
    "dineroenimagen.com",
    "bbc.com",
    "reuters.com",
    "france24.com",
}

TEMAS = {
    "Plan México": [
        '“Plan México”', 'Plan Mexico', 'Polos de Bienestar',
        'relocalización industrial México', 'corredor interoceánico', 'nearshoring México'
    ],
    "Banca de Desarrollo / Bancomext / Nafin": [
        'banca de desarrollo México', 'Nafin', 'Bancomext',
        'FIRA', 'Financiera Nacional de Desarrollo', 'Nacional Financiera'
    ],
    "Sectores Productivos y Económicos": [
        'sectores productivos México', 'industria manufacturera México',
        'inversión productiva México', 'exportaciones México', 'nearshoring México'
    ],
    "Mercados Financieros": [
        'Bolsa Mexicana de Valores', 'BMV', 'peso mexicano', 'tipo de cambio dólar',
        'tasa de interés Banxico', 'bonos M', 'mercados financieros México'
    ],
    "Inversión": [
        'inversión productiva México', 'inversión pública México',
        'IED México', 'inversión extranjera directa México'
    ],
    "Aranceles y Comercio": [
        'aranceles México', 'impuesto a importaciones México',
        'medidas comerciales México', 'comercio exterior México'
    ],
    "Indicadores Económicos Relevantes": [
        'PIB México', 'inflación México', 'IGAE', 'Indicador Oportuno de la Actividad Económica',
        'PMI México', 'IMSS empleo', 'tasa de desempleo México', 'INEGI'
    ],
}

HL = "es-419"; GL = "MX"; CEID = "MX:es-419"
UA = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X) NewsBot/1.3"}
MESES_ES = {1:"enero",2:"febrero",3:"marzo",4:"abril",5:"mayo",6:"junio",
            7:"julio",8:"agosto",9:"septiembre",10:"octubre",
            11:"noviembre",12:"diciembre"}

# =========================
# FECHA/HORA
# =========================
def ahora_tz():
    return datetime.now(pytz.timezone(TZ))

def hoy_local_date():
    return ahora_tz().date()

def ahora_str():
    return ahora_tz().strftime("%Y-%m-%d (%A) %H:%M %Z")

def fecha_humana():
    dt = ahora_tz()
    return f"{int(dt.strftime('%d'))} de {MESES_ES[int(dt.strftime('%m'))]} de {dt.strftime('%Y')}"

def etiqueta_corte_auto():
    h = ahora_tz().hour
    return "Corte Matutino" if h < 13 else "Corte Vespertino"

# =========================
# UTILIDADES
# =========================
def construir_query(terminos, solo_preferidas=True):
    """
    Usamos when:1d para traer material y luego filtramos por 'hoy' local (CDMX).
    """
    sitios = ""
    if solo_preferidas:
        sitios = " OR ".join([f"site:{s}" for s in FUENTES_PREFERIDAS])
        sitios = f" ({sitios})"
    return f"({' OR '.join(terminos)}){sitios} when:1d"

def google_news_rss(query):
    return f"https://news.google.com/rss/search?q={quote_plus(query)}&hl={HL}&gl={GL}&ceid={CEID}"

def _unpack_google_news_link(link):
    """
    Si el link es news.google.com, intenta extraer ?url=... o seguir redirecciones (HEAD).
    """
    try:
        u = urlparse(link)
        if "news.google.com" in (u.netloc or ""):
            qs = parse_qs(u.query)
            if "url" in qs and qs["url"]:
                return qs["url"][0]
        # Si no hay ?url=, intentamos HEAD
        try:
            r = requests.head(link, headers=UA, timeout=10, allow_redirects=True)
            if r.ok:
                return r.url
        except Exception:
            pass
    except Exception:
        pass
    return link

def dominio(link):
    try:
        real = _unpack_google_news_link(link)
        m = re.search(r"https?://([^/]+)/?", real)
        if not m: return ""
        dom = m.group(1).lower().replace("www.", "")
        if dom.endswith(":443"):
            dom = dom[:-4]
        return dom
    except Exception:
        return ""

def _parsed_datetime_local(entry):
    """
    Devuelve datetime local (CDMX) usando published_parsed/updated_parsed
    y con fallback a 'published'/'updated' string si es posible.
    """
    tz = pytz.timezone(TZ)
    t = entry.get("published_parsed") or entry.get("updated_parsed")
    if t:
        return datetime.fromtimestamp(time.mktime(t), tz=pytz.utc).astimezone(tz)
    for key in ("published", "updated"):
        s = entry.get(key)
        if s:
            try:
                from email.utils import parsedate_to_datetime
                dt = parsedate_to_datetime(s)
                if dt.tzinfo is None:
                    dt = pytz.utc.localize(dt)
                return dt.astimezone(tz)
            except Exception:
                pass
    return None

def is_published_today(entry):
    dt_local = _parsed_datetime_local(entry)
    return bool(dt_local and dt_local.date() == hoy_local_date())

def fetch_feed(url, reintentos=REINTENTOS, espera=ESPERA_SEG):
    last_exc = None
    for i in range(reintentos):
        try:
            resp = requests.get(url, headers=UA, timeout=HTTP_TIMEOUT)
            resp.raise_for_status()
            feed = feedparser.parse(resp.content)
            if getattr(feed, "entries", None):
                return feed
        except Exception as e:
            last_exc = e
        time.sleep(espera * (i+1))
    print(f"[WARN] RSS vacío/falló: {url} ({last_exc})")
    return None

# =========================
# RECOLECTA (HOY) + PRIORIDAD FUENTES
# =========================
def recolectar_hoy_por_tema(solo_preferidas=True):
    """
    Cuando solo_preferidas=True, limita por FUENTES_PREFERIDAS.
    Si no hay resultados, vuelve a llamarse con solo_preferidas=False.
    """
    todo = []
    for tema, terminos in TEMAS.items():
        q = construir_query(terminos, solo_preferidas=solo_preferidas)
        feed = fetch_feed(google_news_rss(q))
        if not feed:
            continue
        vistos = set(); capturados = []
        # Debug opcional:
        # print(f"[DBG] tema={tema} feed={len(feed.entries)} hoy={sum(is_published_today(e) for e in feed.entries)}")
        for entry in feed.entries:
            if not is_published_today(entry):
                continue
            link = entry.get("link", "")
            link_real = _unpack_google_news_link(link)
            dom = dominio(link_real)

            if solo_preferidas and dom not in FUENTES_PREFERIDAS:
                continue

            titulo = entry.get("title", "").strip()
            sumry = re.sub(r"\s+", " ", re.sub(r"<[^>]+>", "", entry.get("summary","")))
            key = (titulo.lower(), dom.lower())
            if key in vistos:
                continue
            vistos.add(key)
            capturados.append({
                "tema": tema,
                "titulo": titulo,
                "link": link_real,     # link real del medio
                "dominio": dom,
                "summary_feed": sumry,
                "fecha_local": ahora_str()
            })
            if len(capturados) >= MAX_POR_TEMA:
                break
        if capturados:
            todo.extend(capturados)
    return todo

# =========================
# EXTRACCIÓN / SÍNTESIS
# =========================
def _extract_with_newspaper(url):
    art = Article(url, language='es')
    art.download()
    art.parse()
    return art.title or "", art.text or ""

def _extract_with_readability(url):
    r = requests.get(url, headers=UA, timeout=HTTP_TIMEOUT)
    r.raise_for_status()
    doc = ReadabilityDoc(r.text)
    title = doc.short_title()
    html = doc.summary()
    txt = re.sub(r"<[^>]+>", "", html)
    txt = re.sub(r"\s+", " ", txt).strip()
    return title, txt

def extraer_texto(url):
    last_exc = None
    if _HAVE_NEWSPAPER:
        try:
            return _extract_with_newspaper(url)
        except Exception as e:
            last_exc = e
    if _HAVE_READABILITY:
        try:
            return _extract_with_readability(url)
        except Exception as e:
            last_exc = e
    raise RuntimeError(f"No se pudo extraer texto ({last_exc})")

def limpiar_parrafos(texto):
    if not texto:
        return []
    texto = texto.replace("\r", "\n")
    texto = re.sub(r"\n{2,}", "\n\n", texto)
    partes = [p.strip() for p in re.split(r"\n\s*\n|(?<=\.)\s{2,}", texto) if p.strip()]
    parrafos = []
    for p in partes:
        if len(p) < 80:
            continue
        parrafos.append(p)
    return parrafos

def sintetizar_en_parrafos(texto, max_parrafos=4, min_parrafos=3):
    pars = limpiar_parrafos(texto)
    if not pars:
        oraciones = re.split(r"(?<=[\.\!\?])\s+", texto)
        oraciones = [o.strip() for o in oraciones if len(o.strip()) > 40]
        bloque = []; parrafos = []
        for o in oraciones:
            bloque.append(o)
            if len(" ".join(bloque)) > 350:
                parrafos.append(" ".join(bloque)); bloque = []
        if bloque:
            parrafos.append(" ".join(bloque))
        pars = parrafos
    if len(pars) < min_parrafos:
        return pars
    return pars[:max_parrafos]

def fallback_parrafos_desde_feed(summary_feed):
    if not summary_feed:
        return []
    oraciones = re.split(r"(?<=[\.\!\?])\s+", summary_feed)
    oraciones = [o.strip() for o in oraciones if len(o.strip()) > 30]
    if not oraciones:
        return [summary_feed]
    parrafos, bloque = [], []
    for o in oraciones:
        bloque.append(o)
        if len(" ".join(bloque)) >= 350:
            parrafos.append(" ".join(bloque)); bloque = []
    if bloque:
        parrafos.append(" ".join(bloque))
    return parrafos[:4] if len(parrafos) > 4 else parrafos

def anexar_resumenes(items):
    """
    Añade 'resumen_parrafos' para cada item.
    Si no se puede extraer cuerpo, usa summary del feed como plan B.
    Si ninguno funciona, excluye la nota.
    """
    final = []
    for it in items:
        pars = []
        try:
            _title, texto = extraer_texto(it["link"])
            pars = sintetizar_en_parrafos(texto, max_parrafos=4, min_parrafos=3)
        except Exception as e:
            print(f"[INFO] No se pudo extraer {it.get('link')}: {e}. Uso resumen del feed.")
            pars = fallback_parrafos_desde_feed(it.get("summary_feed", ""))

        if not pars:
            continue

        it2 = dict(it)
        it2["resumen_parrafos"] = pars
        final.append(it2)
    return final

# =========================
# SALIDAS (Markdown / CSV / DOCX)
# =========================
def armar_markdown(fecha_str, items):
    lines = [f"# Apertura de noticias — {fecha_str}\n",
             "_Titulares automáticos (solo HOY; prioriza fuentes; si no hay, abre a otras)._", ""]
    df = pd.DataFrame(items)
    if df.empty:
        lines.append("> No se encontraron noticias hoy para los temas seleccionados.")
        return "\n".join(lines)
    for tema in TEMAS.keys():
        sub = df[df["tema"] == tema]
        if sub.empty:
            continue
        lines.append(f"## {tema}")
        for _, r in sub.iterrows():
            lines.append(f"- **{r['titulo']}** _(Fuente: {r['dominio']})_ — [Leer]({r['link']})")
        lines.append("")
    return "\n".join(lines)

def guardar_salida_diaria(texto, out_root=OUT_DIR_ROOT, nombre_base=NOMBRE_BASE):
    os.makedirs(out_root, exist_ok=True)
    ahora = ahora_tz()
    carpeta_dia = os.path.join(out_root, ahora.strftime("%Y-%m-%d"))
    os.makedirs(carpeta_dia, exist_ok=True)
    dia = int(ahora.strftime("%d"))
    mes = MESES_ES[int(ahora.strftime("%m"))]
    nombre_archivo = f"{nombre_base}_{dia}_{mes}.md"
    ruta = os.path.join(carpeta_dia, nombre_archivo)
    with open(ruta, "w", encoding="utf-8") as f:
        f.write(texto)
    print(f"[OK] Archivo diario guardado en: {ruta}")
    return ruta, carpeta_dia

def anexar_log(items, out_root=OUT_DIR_ROOT, nombre=NOMBRE_CSV_LOG):
    if not items: return None
    os.makedirs(out_root, exist_ok=True)
    ruta = os.path.join(out_root, nombre)
    df = pd.DataFrame(items)[["tema","titulo","link","dominio","fecha_local"]]
    if os.path.exists(ruta):
        cur = pd.read_csv(ruta)
        pd.concat([cur, df], ignore_index=True).drop_duplicates(
            subset=["tema","titulo","link"], keep="first"
        ).to_csv(ruta, index=False)
    else:
        df.to_csv(ruta, index=False)
    return ruta

def construir_docx(fecha_hum, items, carpeta_dia, nombre_base=NOMBRE_BASE):
    doc = Document()
    titulo = doc.add_paragraph(f"Noticias económicas relevantes - {fecha_hum}")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = titulo.runs[0]; r0.bold = True; r0.font.size = Pt(16)
    doc.add_paragraph("")

    if not items:
        p = doc.add_paragraph("No se encontraron noticias hoy para los temas seleccionados.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        df = pd.DataFrame(items)
        for tema in TEMAS.keys():
            sub = df[df["tema"] == tema]
            if sub.empty:
                continue

            h = doc.add_paragraph(tema)
            h.runs[0].bold = True
            h.runs[0].font.size = Pt(13)

            for _, r in sub.iterrows():
                p1 = doc.add_paragraph()
                run1 = p1.add_run("Encabezado: "); run1.bold = True
                p1.add_run(r["titulo"])

                p2lead = doc.add_paragraph()
                lead = p2lead.add_run("Síntesis de la noticia: "); lead.bold = True

                pars = r.get("resumen_parrafos", [])
                if not pars:
                    doc.add_paragraph("(No fue posible generar una síntesis confiable.)")
                else:
                    for par in pars:
                        doc.add_paragraph(par)

                p3 = doc.add_paragraph()
                run3 = p3.add_run("Enlace a la nota: "); run3.bold = True
                p3.add_run(r["link"])

                p4 = doc.add_paragraph(f"Fuente: {r['dominio']}")
                p4.runs[0].italic = True

                doc.add_paragraph("")

    dia = int(ahora_tz().strftime("%d"))
    mes = MESES_ES[int(ahora_tz().strftime("%m"))]
    nombre_docx = f"{nombre_base}_{dia}_{mes}.docx"
    ruta_docx = os.path.join(carpeta_dia, nombre_docx)
    doc.save(ruta_docx)
    print(f"[OK] DOCX guardado en: {ruta_docx}")
    return ruta_docx

# ---------- Email ----------
def enviar_correo(md_path, fecha_str, docx_path=None, items_count=0):
    if items_count <= 0:
        print("[INFO] No hay noticias válidas hoy. No se envía correo.")
        return False
    if not GMAIL_USER or not GMAIL_APP_PASS:
        print("[WARN] GMAIL_USER/GMAIL_APP_PASS no están definidos. No se envía correo.")
        return False

    with open(md_path, "r", encoding="utf-8") as f:
        contenido_md = f.read()

    label = CORTE_LABEL or etiqueta_corte_auto()
    msg = MIMEMultipart()
    msg["From"] = GMAIL_USER
    msg["To"] = ", ".join(EMAIL_TO)
    if EMAIL_CC:
        msg["Cc"] = ", ".join(EMAIL_CC)
    msg["Subject"] = f"{ASUNTO_BASE} — {label} — {fecha_str}"

    cuerpo = f"""Hola, Mr. PC:

Apertura generada automáticamente (solo HOY; prioriza fuentes; si no hay, abre a otras).

Total notas: {items_count}
"""
    msg.attach(MIMEText(cuerpo + "\n\n" + contenido_md, "plain", "utf-8"))

    part_md = MIMEBase("application", "octet-stream")
    with open(md_path, "rb") as f:
        part_md.set_payload(f.read())
    encoders.encode_base64(part_md)
    part_md.add_header("Content-Disposition", f'attachment; filename="{os.path.basename(md_path)}"')
    msg.attach(part_md)

    if docx_path and os.path.exists(docx_path):
        part_docx = MIMEBase("application", "vnd.openxmlformats-officedocument.wordprocessingml.document")
        with open(docx_path, "rb") as f:
            part_docx.set_payload(f.read())
        encoders.encode_base64(part_docx)
        part_docx.add_header("Content-Disposition", f'attachment; filename="{os.path.basename(docx_path)}"')
        msg.attach(part_docx)

    dests = EMAIL_TO + EMAIL_CC
    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as s:
        s.login(GMAIL_USER, GMAIL_APP_PASS)
        s.sendmail(GMAIL_USER, dests, msg.as_string())
    print("[OK] Correo enviado a:", dests)
    return True

# =========================
# MAIN
# =========================
if __name__ == "__main__":
    fecha_trace = ahora_str()
    fecha_hum = fecha_humana()

    # 1) Intento 1: SOLO FUENTES PREFERIDAS (de HOY, con desempaquetado de link)
    items_hoy = recolectar_hoy_por_tema(solo_preferidas=True)

    # 2) Intento 2: si quedó vacío, abre a CUALQUIER FUENTE (de HOY)
    if not items_hoy:
        print("[INFO] No hubo resultados en fuentes preferidas; buscando en otras fuentes (HOY).")
        items_hoy = recolectar_hoy_por_tema(solo_preferidas=False)

    # 3) Síntesis (3–4 párrafos). Si no hay cuerpo, usa resumen del feed.
    items_con_resumen = anexar_resumenes(items_hoy)

    # 4) Markdown y guardado
    md = armar_markdown(fecha_trace, items_con_resumen)
    md_path, carpeta_dia = guardar_salida_diaria(md)

    # 5) DOCX
    docx_path = construir_docx(fecha_hum, items_con_resumen, carpeta_dia)

    # 6) Log
    _ = anexar_log(items_con_resumen)

    # 7) Marca OK
    Path(carpeta_dia, ".ok").write_text("ok", encoding="utf-8")

    # 8) Envío (solo si hay ≥1 nota)
    enviar_correo(md_path, fecha_trace, docx_path=docx_path, items_count=len(items_con_resumen))